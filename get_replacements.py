import re
from docx import Document
from database import DataBase
import os

class Replacement:
    def __init__(self) -> None:
        self.document = Document(os.getcwd() + r"\documents\zameni.docx")

    def update_replacement(self):
        self.document = Document(os.getcwd() + r"\documents\zameni.docx")

    def zameni_group(self, user_id):

        group = DataBase().get_group(user_id)
        aboba = []

        for paragraph in self.document.paragraphs:
            
            if "на " in paragraph.text:
                aboba.append(paragraph.text)
        # Перебор
        for table in enumerate(self.document.tables, start=0):
            for x in range(0, 3):
                row = table[1].column_cells(x)
                
                gpz = [re.compile(r'\s{2,}').sub(' ', i.text.strip()) for i in row]
                for i in range(0, len(gpz), 2):
                    if gpz[i] == group:
                        if table[0] == 0 or table[0] == 1:
                            
                            date = aboba[0]
                            for table2 in enumerate(self.document.tables[2:-1], start=0):
                                for y in range(0, 3):
                                    row2 = table2[1].column_cells(y)
                                    gpz2 = [re.compile(r'\s{2,}').sub(' ', y.text.strip()) for y in row2]

                                    for t in range(0, len(gpz2), 2):
                                        if gpz2[t] == group:
                                            return(str(f'Замены для группы {gpz[i]} {date}\n{str(gpz[i+1])}\n\nЗамены для группы {aboba[-1]}:\n{gpz2[t+1]}'))
                            return(str(f'Замены для группы {gpz[i]} {date}\n{str(gpz[i+1])}'))

                        elif table[0] == 2 or table[0] == 3:
                            date = aboba[2]
                            return(str(f'Замены для группы {gpz[i]} {date}\n{str(gpz[i+1])}'))

        return "Замен нет"